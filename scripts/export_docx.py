#!/usr/bin/env python3
"""Export alignment memo and revision packet to Capella-ready .docx files."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path("/workspace")
DOCS = ROOT / "docs"
OUT = ROOT / "downloads"
ARTIFACTS = Path("/opt/cursor/artifacts")


def set_run_font(run, *, bold=None, italic=None, size=12, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_runs(paragraph, text, *, size=12, bold=False, italic=False, color=None):
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    parts = pattern.split(text)
    if not parts:
        run = paragraph.add_run("")
        set_run_font(run, size=size, bold=bold, italic=italic, color=color)
        return
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True, italic=italic, color=color)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=size, italic=True, color=color)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=size, italic=True, color=color)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, bold=bold, italic=italic, color=color)


def style_paragraph(p, *, after=6, before=0, align=None):
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.line_spacing = 2.0
    if align:
        p.alignment = align
    return p


NEW_COLOR = RGBColor(0x9A, 0x34, 0x12)


def is_new_heading(text: str) -> bool:
    return "NEW" in text.upper()


def add_heading_styled(doc, text, level):
    sizes = {1: 16, 2: 14, 3: 12, 4: 12}
    p = doc.add_paragraph()
    style_paragraph(p, after=8, before=12)
    color = NEW_COLOR if is_new_heading(text) else RGBColor(0x1F, 0x3A, 0x5F)
    add_runs(p, text, size=sizes.get(level, 12), bold=True, color=color)
    return p


def add_highlight_box(doc, text):
    p = doc.add_paragraph()
    style_paragraph(p, after=8, before=8)
    add_runs(
        p,
        "HIGHLIGHT — NEW EFFORT: " + text,
        size=12,
        bold=True,
        color=NEW_COLOR,
    )


def split_row(line: str) -> list[str]:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def is_separator(line: str) -> bool:
    return bool(re.match(r"^\s*\|?\s*:?-{3,}", line))


def add_table(doc, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j in range(cols):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.line_spacing = 1.0
            value = row[j] if j < len(row) else ""
            add_runs(p, value, size=10, bold=(i == 0))
    doc.add_paragraph()


def convert_md(md_path: Path, title: str, subtitle: str, out_path: Path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("Walker — RSCH-V8927 / BMGT8044")
    set_run_font(run, size=10, italic=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Download copy for Project Plan use  ·  Page ")
    set_run_font(run, size=10)
    # PAGE field
    run2 = footer.add_run()
    fld = run2._element
    from docx.oxml import OxmlElement

    def add_page_field(paragraph):
        run = paragraph.add_run()
        r = run._r
        fc1 = OxmlElement("w:fldChar")
        fc1.set(qn("w:fldCharType"), "begin")
        r.append(fc1)
        run_instr = paragraph.add_run()
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        run_instr._r.append(instr)
        run_end = paragraph.add_run()
        fc2 = OxmlElement("w:fldChar")
        fc2.set(qn("w:fldCharType"), "end")
        run_end._r.append(fc2)
        set_run_font(run, size=10)
        set_run_font(run_instr, size=10)
        set_run_font(run_end, size=10)

    add_page_field(footer)

    p = doc.add_paragraph()
    style_paragraph(p, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_runs(p, "Max D. Walker", size=12, bold=True)
    p = doc.add_paragraph()
    style_paragraph(p, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_runs(p, "Capella University", size=12)
    p = doc.add_paragraph()
    style_paragraph(p, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_runs(p, "RSCH-V8927 Doctoral Project Development — Framework Development", size=12)
    p = doc.add_paragraph()
    style_paragraph(p, after=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_runs(p, subtitle, size=12, italic=True)

    p = doc.add_paragraph()
    style_paragraph(p, after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_runs(p, title, size=16, bold=True, color=RGBColor(0x1F, 0x3A, 0x5F))

    p = doc.add_paragraph()
    style_paragraph(p, after=12)
    add_runs(
        p,
        "Items marked NEW or HIGHLIGHT are new written efforts that were not present as steps in the prior data-collection draft. Paste Section 9 of the revision packet into the Project Plan. Use the interview protocol and artifact pipeline as appendices.",
        size=12,
        italic=True,
    )

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    # skip first markdown H1 (replaced by title page)
    if lines and lines[0].startswith("# "):
        i = 1
        while i < len(lines) and not lines[i].strip():
            i += 1

    table_buf: list[list[str]] = []

    def flush_table():
        nonlocal table_buf
        if table_buf:
            # drop separator row if present
            cleaned = [r for r in table_buf if not all(re.match(r"^:?-+:?$", c or "") for c in r)]
            add_table(doc, cleaned)
            table_buf = []

    while i < len(lines):
        line = lines[i]
        raw = line.rstrip()
        if raw.startswith("|") and "|" in raw[1:]:
            if is_separator(raw):
                i += 1
                continue
            table_buf.append(split_row(raw))
            i += 1
            continue
        else:
            flush_table()

        if not raw.strip():
            i += 1
            continue
        if raw.startswith("```"):
            # skip mermaid/code fences
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                i += 1
            i += 1
            continue
        if raw == "---":
            i += 1
            continue
        if raw.startswith("#### "):
            add_heading_styled(doc, raw[5:].strip(), 4)
        elif raw.startswith("### "):
            add_heading_styled(doc, raw[4:].strip(), 3)
        elif raw.startswith("## "):
            add_heading_styled(doc, raw[3:].strip(), 2)
        elif raw.startswith("# "):
            add_heading_styled(doc, raw[2:].strip(), 1)
        elif raw.startswith("> "):
            p = doc.add_paragraph()
            style_paragraph(p, after=8, before=4)
            p.paragraph_format.left_indent = Inches(0.5)
            add_runs(p, raw[2:].strip(), size=12, italic=True)
        elif re.match(r"^\d+\.\s", raw):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.line_spacing = 2.0
            # clear default
            if p.runs:
                p.runs[0].text = ""
            add_runs(p, re.sub(r"^\d+\.\s", "", raw), size=12)
        elif raw.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.line_spacing = 2.0
            add_runs(p, raw[2:], size=12)
        else:
            p = doc.add_paragraph()
            style_paragraph(p)
            add_runs(p, raw, size=12)
        i += 1

    flush_table()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    packet = convert_md(
        DOCS / "collection-analysis-revision-packet.md",
        "Collection-Plus-Analysis Revision Packet",
        "Paste-ready protocol, artifact pipeline, analysis bridge, citations, language, and ethics",
        OUT / "WALKER_Collection_Analysis_Revision_Packet.docx",
    )
    assessment = convert_md(
        DOCS / "alignment-assessment-data-collection.md",
        "Alignment Assessment of Data Collection Against the Gap",
        "Scored against The Alignment Map: Guiding Questions",
        OUT / "WALKER_Alignment_Assessment_Data_Collection.docx",
    )
    print(packet)
    print(assessment)
    if ARTIFACTS.exists():
        for src in (packet, assessment):
            dest = ARTIFACTS / src.name
            dest.write_bytes(src.read_bytes())
            print("artifact", dest)


if __name__ == "__main__":
    main()
